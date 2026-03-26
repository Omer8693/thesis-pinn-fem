"""
generate_word.py
Generates a comprehensive Word (.docx) project report for the NAS-MCO-PINNs master's thesis.
Run:  python reports/professor_brief/generate_word.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# PATHS
# ---------------------------------------------------------------------------
BASE   = os.path.dirname(os.path.abspath(__file__))
ROOT   = os.path.normpath(os.path.join(BASE, "..", ".."))
L1     = os.path.join(ROOT, "level1_single_shot",    "results")
L2     = os.path.join(ROOT, "level2_timestepper",    "results")
L3     = os.path.join(ROOT, "level3_hybrid_fem",     "results")
L4     = os.path.join(ROOT, "level4_distortion",     "results")
L5     = os.path.join(ROOT, "level5_refinement",     "results")
L8     = os.path.join(ROOT, "level8_nas_mco_pinn",   "results")
L9     = os.path.join(ROOT, "level9_sa_fourier",     "results", "plots")
OUT    = os.path.join(BASE, "NAS_PINNS3_professor_brief.docx")

# ---------------------------------------------------------------------------
# COLOR CONSTANTS
# ---------------------------------------------------------------------------
NAVY       = RGBColor(0x1a, 0x23, 0x7e)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = "D6E4F0"
TABLE_NAV  = "1a237e"
GRAY_BG    = "F2F2F2"

# Full reference string used throughout the document
FEM_REF = (
    "D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen (2026), "
    "\"Mitigating distortions in cast automotive subframes: A finite element simulation "
    "approach,\" International Journal of Advanced Manufacturing Technology. "
    "DOI: 10.1007/s00170-026-17515-w"
)
FEM_SHORT = "D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen (2026)"

# ---------------------------------------------------------------------------
# LOW-LEVEL XML HELPERS
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_para_shading(para, hex_color):
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# DOCUMENT SETUP
# ---------------------------------------------------------------------------

def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.54))
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    return doc


# ---------------------------------------------------------------------------
# HELPER FUNCTIONS
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.name  = "Times New Roman"
        run.font.size  = Pt(13)
        run.font.color.rgb = NAVY
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(6)
    else:
        run.font.name  = "Times New Roman"
        run.font.size  = Pt(11)
        run.font.color.rgb = DARK_GRAY
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(4)
    return para


def add_para(doc, text, bold_prefix=None, indent=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    if indent:
        para.paragraph_format.left_indent = Cm(1.0)
    if bold_prefix:
        br = para.add_run(bold_prefix + " ")
        br.bold       = True
        br.font.name  = "Times New Roman"
        br.font.size  = Pt(11)
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    else:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    return para


def add_image(doc, path, width_cm=14, caption=None):
    if not os.path.isfile(path):
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run()
        run.add_picture(path, width=Cm(width_cm))
    except Exception:
        return
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after  = Pt(10)
        for run in cap.runs:
            run.font.name   = "Times New Roman"
            run.font.size   = Pt(10)
            run.font.italic = True


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, TABLE_NAV)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold           = True
        run.font.color.rgb = WHITE
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(10)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = LIGHT_BLUE if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    set_cell_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


def add_caption(doc, text):
    cap = doc.add_paragraph(text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(10)
    for r in cap.runs:
        r.font.name   = "Times New Roman"
        r.font.size   = Pt(10)
        r.font.italic = True


def add_equation(doc, eq_text, note=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(1.5)
    para.paragraph_format.right_indent = Cm(1.5)
    set_para_shading(para, GRAY_BG)
    run = para.add_run(eq_text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.bold      = True
    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_before = Pt(2)
        np_.paragraph_format.space_after  = Pt(8)
        np_.paragraph_format.left_indent  = Cm(1.5)
        nr = np_.add_run(note)
        nr.font.name   = "Times New Roman"
        nr.font.size   = Pt(10)
        nr.font.italic = True
    return para


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr    = para._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1a237e")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    run  = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.5 * level)


# ---------------------------------------------------------------------------
# MAIN DOCUMENT BUILD
# ---------------------------------------------------------------------------

def build_document():
    doc = create_document()

    # ========================================================================
    # TITLE BLOCK
    # ========================================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after  = Pt(6)
    tr = title_para.add_run(
        "NAS-PINNs: Neural Architecture Search for FEM Time-Step Acceleration "
        "in Industrial Thermal Simulation"
    )
    tr.bold           = True
    tr.font.name      = "Times New Roman"
    tr.font.size      = Pt(14)
    tr.font.color.rgb = NAVY

    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.paragraph_format.space_after = Pt(2)
    for text, bold in [("Omer Cetinkaya", True), ("  |  IKT590-G 26V — Master's Thesis  |  March 2026", False)]:
        r = auth.add_run(text)
        r.bold       = bold
        r.font.name  = "Times New Roman"
        r.font.size  = Pt(11)

    add_horizontal_rule(doc)

    # ========================================================================
    # SECTION 1: WHAT IS THIS PROJECT?
    # ========================================================================
    add_heading(doc, "1. What Is This Project?", level=1)
    add_para(doc,
        "This project investigates whether Physics-Informed Neural Networks (PINNs) can replace "
        "intermediate Finite Element Method (FEM) time steps in a transient thermal simulation. "
        "The physical setting is water quenching of an A356 aluminium alloy automotive subframe "
        "— a real manufacturing process in automotive production."
    )
    add_para(doc,
        "When a hot aluminium casting (at 540 °C) is plunged into cold water (20 °C), the surface "
        "cools much faster than the core. This non-uniform temperature field creates thermal "
        "gradients that produce residual stresses, which in turn cause the part to distort. "
        "Engineers need to predict this distortion before manufacturing the part, and they use FEM "
        "to do so. The problem is that FEM is slow: each full simulation takes about 184 seconds, "
        "and when optimizing a part design you may need to run thousands of simulations."
    )
    add_para(doc,
        "Our hypothesis was: can we train a neural network once, and then use it to predict the "
        "temperature field at any point in time — replacing most or all of the FEM time steps — "
        "so that subsequent predictions take milliseconds instead of minutes?"
    )
    add_para(doc,
        "We answered this question through a nine-level framework. Each level added one new "
        "technical contribution. The best result, at Level 9, achieves an L2 error of 0.014 "
        "with an inference time of approximately 0.01 seconds — an 18,400-fold speedup compared "
        "to FEM."
    )

    # ========================================================================
    # SECTION 2: THE PHYSICAL PROBLEM
    # ========================================================================
    add_heading(doc, "2. The Physical Problem", level=1)

    add_heading(doc, "2.1 Governing Equation", level=2)
    add_para(doc,
        "The temperature field T(t, x, y) inside the aluminium part evolves according to the "
        "transient heat conduction equation:"
    )
    add_equation(doc,
        "rho * cp * dT/dt  =  k * (d^2T/dx^2 + d^2T/dy^2)",
        note=(
            "rho = 2670 kg/m^3 (density),  cp = 963 J/(kg*K) (specific heat),  "
            "k = 151 W/(m*K) (thermal conductivity) for A356 aluminium."
        )
    )
    add_para(doc,
        "On the outer surface of the part, the water extracts heat by convection. We model this "
        "with a Robin (mixed) boundary condition:"
    )
    add_equation(doc,
        "-k * dT/dn  =  h(T) * (T - T_water)",
        note=(
            "n is the outward surface normal. h(T) is the heat transfer coefficient "
            "[W/m^2*K] — a nonlinear function of surface temperature, modelling boiling regime "
            "transitions. T_water = 20 deg C."
        )
    )
    add_para(doc,
        "The initial condition is T(x, y, 0) = 540 °C (uniform temperature at quench start). "
        "The 2D computational domain is a rectangular cross-section of 1.3 m × 0.6 m. "
        "The simulation covers 30 seconds, discretised into 21 FEM time steps."
    )

    add_heading(doc, "2.2 FEM Reference Data — NOT Implemented by Us", level=2)
    add_para(doc,
        "We did not implement FEM ourselves. All finite element simulation data used as "
        "ground truth in this project comes from the following published paper:",
        bold_prefix="Important:"
    )
    add_para(doc, FEM_REF, indent=True)
    add_para(doc,
        "From this paper we extracted three types of data:"
    )
    add_bullet(doc,
        "A356 material flow stress parameters (F, n, m as functions of temperature) from "
        "Table 1 of D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen (2026). "
        "These were used to build our material interpolation module in src/physics_model.py."
    )
    add_bullet(doc,
        "The heat transfer coefficient curve h(T), which transitions from film boiling at high "
        "temperatures to nucleate boiling and then forced convection as the surface cools. "
        "We approximated this with a piecewise function."
    )
    add_bullet(doc,
        "Reference temperature histories at multiple points in the casting, and the FEM-predicted "
        "distortion at 8 CMM measurement points. These served as our ground truth for computing "
        "L2 error and mean absolute error (MAE)."
    )

    add_table(doc,
        headers=["Property", "Symbol", "Value", "Unit"],
        rows=[
            ["Density",                  "rho",    "2670",  "kg/m^3"],
            ["Specific heat",            "cp",     "963",   "J/(kg*K)"],
            ["Thermal conductivity",     "k",      "151",   "W/(m*K)"],
            ["Initial temperature",      "T_0",    "540",   "deg C"],
            ["Water temperature",        "T_inf",  "20",    "deg C"],
            ["Convection coeff. (ref)",  "h_ref",  "5000",  "W/(m^2*K)"],
            ["Simulation duration",      "t_end",  "30",    "s"],
            ["Number of FEM time steps", "N_t",    "21",    "—"],
        ],
        col_widths=[5.0, 2.0, 2.5, 3.0]
    )
    add_caption(doc,
        "Table 1. A356 aluminium alloy properties used in this project. "
        "Source: D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen (2026), Table 1."
    )

    # Baseline figures
    add_image(doc,
        os.path.join(L1, "baseline", "baseline_a356_params.png"),
        width_cm=14,
        caption=(
            "Figure 1. A356 flow stress parameters (F, n, m) as functions of temperature, "
            "digitised from Table 1 of D. Mortensen, G. Noorsumar, H.G. Fjaer, "
            "R. Babaei, P.E. Dronen (2026). Used as material input to our PINN."
        )
    )
    add_image(doc,
        os.path.join(L1, "baseline", "baseline_fig17_ht.png"),
        width_cm=13,
        caption=(
            "Figure 2. FEM reference temperature history at a representative point in the "
            "casting (from Figure 17, D. Mortensen, G. Noorsumar, H.G. Fjaer, "
            "R. Babaei, P.E. Dronen (2026)). This curve is our primary validation target."
        )
    )

    # ========================================================================
    # SECTION 3: WHAT IS A PINN?
    # ========================================================================
    add_heading(doc, "3. What Is a Physics-Informed Neural Network?", level=1)
    add_para(doc,
        "A Physics-Informed Neural Network (PINN) is a neural network that is trained not only "
        "to fit data, but also to satisfy the governing physical equations. The idea was "
        "introduced by Raissi, Perdikaris and Karniadakis [1] in 2019."
    )
    add_para(doc,
        "In our case, the network takes as input the coordinates (t, x, y) and outputs the "
        "predicted temperature T. The network is trained by minimising three types of loss "
        "simultaneously:"
    )
    add_bullet(doc,
        "Physics loss (L_phys): the PDE residual — the network output must satisfy "
        "rho*cp*dT/dt = k*(d^2T/dx^2 + d^2T/dy^2) at randomly sampled collocation points "
        "inside the domain."
    )
    add_bullet(doc,
        "Boundary condition loss (L_bc): the Robin boundary condition must be satisfied at "
        "points on the outer surface."
    )
    add_bullet(doc,
        "Initial condition loss (L_ic): the predicted temperature at t=0 must equal 540 °C "
        "everywhere."
    )
    add_equation(doc,
        "L_total  =  lambda_phys * L_phys  +  lambda_bc * L_bc  +  lambda_ic * L_ic",
        note=(
            "lambda values are the loss weights. At Levels 1-8 these were fixed. "
            "At Level 9 they become learnable parameters (self-adaptive weights)."
        )
    )
    add_para(doc,
        "The advantage of PINNs over classical FEM is that once trained, inference is essentially "
        "free — a single forward pass through the network. The disadvantage is that training is "
        "slow (minutes) and the accuracy depends heavily on the network architecture and training "
        "setup. Our project focused on using Neural Architecture Search (NAS) to automatically "
        "find the best architecture."
    )

    # ========================================================================
    # SECTION 4: NEURAL ARCHITECTURE SEARCH
    # ========================================================================
    add_heading(doc, "4. Neural Architecture Search (NAS)", level=1)
    add_para(doc,
        "Neural Architecture Search automates the process of choosing the right network structure. "
        "Instead of manually trying many combinations of layers, neurons, and activation functions, "
        "we run an automated search that evaluates many candidate architectures and finds the best one."
    )
    add_para(doc,
        "Our search space covers: number of layers (3–6), neurons per layer (64–160, same for all "
        "layers), and activation function (Tanh, GELU, SiLU, ReLU). We compared three NAS strategies:"
    )
    add_bullet(doc,
        "Bayesian Optimisation [8] (via scikit-optimize library): Fits a Gaussian Process "
        "surrogate model on past results and uses it to select the next candidate. Single-objective: "
        "minimises L2 error. Good when the landscape is smooth."
    )
    add_bullet(doc,
        "NSGA-II [6] (via pymoo library, Deb et al. 2002): A multi-objective evolutionary "
        "algorithm. Minimises both L2 error AND number of parameters. Finds compact architectures "
        "that generalise well."
    )
    add_bullet(doc,
        "NSGA-III [7] (via pymoo library, Deb & Jain 2014): Extension of NSGA-II with "
        "structured reference points. Tends to find very compact architectures, sometimes at "
        "the cost of slightly higher error."
    )
    add_para(doc,
        "Each NAS run evaluates approximately 500 candidate architectures. The winning architecture "
        "is then used for the full training run at each level. In 2D, the search takes roughly "
        "400–600 seconds. Our main finding was that NSGA-II consistently found the best trade-off "
        "between accuracy and size in the quenching problem."
    )

    # ========================================================================
    # SECTION 5: NINE-LEVEL FRAMEWORK
    # ========================================================================
    add_heading(doc, "5. The Nine-Level Framework", level=1)
    add_para(doc,
        "We developed the framework progressively, adding one new technical contribution per "
        "level. This section describes each level and what it contributed."
    )

    add_table(doc,
        headers=["Level", "Title", "Best L2", "Key Contribution"],
        rows=[
            ["L1",    "Single-Shot NAS-PINN",  "0.076",         "Baseline PINN with NAS architecture search"],
            ["L2",    "Temporal Skip Operator", "0.108 (s=2)",   "One PINN predicts s FEM steps ahead — 48% reduction"],
            ["L3",    "Hybrid FEM + PINN",      "~0.12",         "Residual-based adaptive FEM/PINN routing"],
            ["L4",    "Thermal Distortion",     "MAE = 14.4 °C", "Temperature field to CMM distortion mapping"],
            ["L5",    "Adam + L-BFGS",          "0.030",         "Two-stage training: Adam then L-BFGS refinement"],
            ["L6",    "Poisson Benchmark",      "~0.018",        "Generalization test on Poisson PDE"],
            ["L7A",   "Temporal Analysis",      "~0.025",        "Systematic time-skip regime study"],
            ["L7B",   "Multi-Domain NAS",       "~0.022",        "NAS across 5 different geometry types"],
            ["L8",    "3D MCO-PINNs",           "MAE = 2.19 °C", "3D geometries, MCO adaptive weights, 48 runs"],
            ["L9 *",  "SA + Fourier (best)",    "0.014",         "Learnable loss weights + Fourier feature embedding"],
        ],
        col_widths=[1.5, 4.0, 3.0, 6.0]
    )
    add_caption(doc, "Table 2. Nine-level progressive framework. * = best overall result.")

    add_heading(doc, "5.1 Level 1 — Single-Shot NAS-PINN (Baseline)", level=2)
    add_para(doc,
        "Level 1 is the baseline. We trained a single PINN to predict the temperature T(t, x, y) "
        "globally over the full 30-second simulation. The NAS search identified the best "
        "architecture. The best L2 error was 0.076, achieved by the Bayesian-optimised architecture."
    )
    add_image(doc,
        os.path.join(L1, "fig1_cooling_curves.png"),
        width_cm=14,
        caption=(
            "Figure 3. Level 1 — Cooling curves: PINN-predicted temperature vs. FEM reference "
            "at multiple points in the casting. Each colour is a different spatial location. "
            "The PINN captures the general trend but has moderate error in the first 5 seconds."
        )
    )
    add_image(doc,
        os.path.join(L1, "fig2_spatial_fields.png"),
        width_cm=14,
        caption=(
            "Figure 4. Level 1 — Spatial temperature fields at several time steps. Left: FEM "
            "reference. Right: PINN prediction. The PINN correctly identifies the spatial pattern "
            "but has smoothed errors near the boundary."
        )
    )

    add_heading(doc, "5.2 Level 2 — Temporal Skip Operator (Our Own Design)", level=2)
    add_para(doc,
        "Standard PINNs learn the full time history from t = 0 to t = 30 s. At Level 2 we "
        "designed a different approach: a single network that predicts s time steps ahead from "
        "any given initial condition. This is the temporal skip operator — our own contribution "
        "to the project."
    )
    add_para(doc, "The formula we designed is:")
    add_equation(doc,
        "T_hat(t + s*dt)  =  T_IC  +  tau * f_theta(x, y, tau, T_IC)",
        note=(
            "tau = t / (s*dt) in [0, 1] is the normalized local time within the window. "
            "T_IC is the temperature field at the start of the window. "
            "At tau=0, the output exactly equals T_IC — initial condition is automatically satisfied. "
            "dt = 30/21 ≈ 1.43 s (one FEM time step)."
        )
    )
    add_para(doc,
        "Why we designed it this way: writing the prediction as T_IC plus a learned correction "
        "guarantees that at the start of any window (tau=0), the prediction equals the initial "
        "temperature exactly, without needing a separate IC loss term. This idea is inspired by "
        "early work on neural network solutions to ODEs/PDEs [10], but the skip operator itself "
        "— applying this to predict multiple steps ahead in a quenching simulation — is our "
        "original contribution."
    )
    add_para(doc,
        "A counterintuitive finding: skip=2 gives lower error (L2=0.108) than skip=1 (L2=0.126). "
        "With skip=1, any small FEM numerical noise in consecutive steps accumulates; with skip=2 "
        "the PINN smooths these oscillations, effectively acting as a denoising step."
    )

    add_table(doc,
        headers=["Skip s", "FEM Calls", "FEM Steps Saved", "L2 Error", "MAE (°C)", "Runtime"],
        rows=[
            ["1 (baseline)",     "21", "0%",   "0.126", "43.6", "184 s"],
            ["2 (recommended)",  "11", "48%",  "0.108", "33.2", "91 s"],
            ["4",                 "6", "71%",  "0.204", "57.5", "46 s"],
            ["6",                 "4", "81%",  "0.294", "93.6", "28 s"],
        ],
        col_widths=[3.2, 2.5, 3.5, 2.5, 2.8, 2.0]
    )
    add_caption(doc, "Table 3. Skip operator results (Level 2, Bayesian architecture).")

    add_image(doc,
        os.path.join(L2, "skip_comparison.png"),
        width_cm=14,
        caption=(
            "Figure 5. Level 2 — Skip operator: L2 error vs. skip value s for all three "
            "NAS optimisers. The skip=2 point outperforms skip=1 for all optimisers, "
            "confirming the counterintuitive noise-smoothing result."
        )
    )
    add_image(doc,
        os.path.join(L2, "fem_vs_pinn_skip.png"),
        width_cm=14,
        caption=(
            "Figure 6. Level 2 — FEM vs. PINN comparison at different skip values. "
            "Cooling curves show the PINN accurately tracks the FEM reference "
            "for skip=2 while using only 11 of 21 FEM calls."
        )
    )

    add_heading(doc, "5.3 Levels 3–7 — Intermediate Developments", level=2)
    add_para(doc,
        "Levels 3–7 each added one capability that contributed to the final result:"
    )
    add_bullet(doc,
        "Level 3 — Hybrid FEM+PINN: An adaptive routing system that uses FEM for time steps "
        "where the PINN residual is high (above a threshold), and the PINN for easier steps. "
        "This reduces total FEM calls without sacrificing accuracy."
    )
    add_bullet(doc,
        "Level 4 — Thermal Distortion: Extended the framework to predict part distortion from "
        "the temperature field. We used a plane-stress FEM model to map temperatures to "
        "displacements and compared with the 8 CMM measurement points from "
        "D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen (2026). "
        "Best MAE = 14.4 °C."
    )
    add_bullet(doc,
        "Level 5 — Two-Stage Training (Adam + L-BFGS): Standard Adam optimiser is good at "
        "finding the basin of attraction of the minimum but struggles to converge tightly. "
        "We follow Adam training with a second stage using L-BFGS, a quasi-Newton method that "
        "converges much more precisely near the minimum. This improved L2 from 0.076 (L1) "
        "to 0.030 at L5."
    )
    add_bullet(doc,
        "Level 6 — Poisson Benchmark: We validated that our NAS-found architectures generalise "
        "beyond the quenching problem by applying them to the steady-state Poisson equation "
        "on a square domain. Achieved L2 ~ 0.018."
    )
    add_bullet(doc,
        "Levels 7A/7B — Temporal and Multi-Domain Analysis: Systematic study of the temporal "
        "skip regime across different training setups (7A), and NAS across five different "
        "geometry types including circle, annulus, L-shape, flower, and square (7B). "
        "These levels showed the NAS approach generalises across geometry."
    )

    add_image(doc,
        os.path.join(L5, "level5_l2_comparison.png"),
        width_cm=14,
        caption=(
            "Figure 7. Level 5 — L2 error comparison across optimisers with two-stage "
            "Adam + L-BFGS training. Bayesian optimisation achieves the best L2 = 0.030 "
            "at this level, a 2.5× improvement over Level 1 (L2 = 0.076)."
        )
    )

    add_heading(doc, "5.4 Level 8 — 3D Multi-Geometry (MCO-PINNs)", level=2)
    add_para(doc,
        "Level 8 extended the framework to three-dimensional domains and four different geometry "
        "shapes: Rectangular Prism, Cylinder, Stacked Cubes, and L-Shape. We ran 48 experiments "
        "(4 domains × 3 NAS optimisers × 4 skip values). We also introduced multi-component "
        "objective (MCO) adaptive loss weighting, which adjusts the physics/BC/IC weights "
        "based on gradient magnitudes during training."
    )
    add_para(doc,
        "The best 3D result was MAE = 2.19 °C on the L-Shape domain with NSGA-II at skip=1. "
        "This is the most 3D-realistic configuration we tested and demonstrates that the "
        "framework is feasible in 3D."
    )
    add_image(doc,
        os.path.join(L8, "fig8_3d_feasibility.png"),
        width_cm=14,
        caption=(
            "Figure 8. Level 8 — 3D feasibility study. Comparison of PINN predicted "
            "temperature fields vs. FEM reference across the four 3D domain shapes "
            "at the z-midplane slice."
        )
    )
    add_image(doc,
        os.path.join(L8, "fig1_thermal_fields.png"),
        width_cm=14,
        caption=(
            "Figure 9. Level 8 — Predicted temperature fields on the z-midplane for all "
            "four 3D domains (Rectangular, Cylinder, Stacked, L-Shape). The PINN correctly "
            "captures the spatial distribution in each case."
        )
    )
    add_image(doc,
        os.path.join(L8, "fig4_summary_table.png"),
        width_cm=14,
        caption=(
            "Figure 10. Level 8 summary table: MAE (°C) for all 48 runs across 4 domains, "
            "3 optimisers (Bayesian, NSGA-II, NSGA-III), and 4 skip values (1, 2, 4, 6)."
        )
    )

    add_table(doc,
        headers=["Domain", "Best Optimizer", "Best Skip", "Best MAE (°C)"],
        rows=[
            ["Rectangular Prism", "NSGA-II",   "1", "3.42"],
            ["Cylinder",          "Bayesian",  "1", "4.81"],
            ["Stacked Cubes",     "NSGA-III",  "2", "3.17"],
            ["L-Shape",           "NSGA-II",   "1", "2.19 *"],
        ],
        col_widths=[4.5, 3.5, 2.5, 4.0]
    )
    add_caption(doc, "Table 4. Level 8 best results per 3D domain. * = best overall 3D result.")

    # ========================================================================
    # SECTION 6: OUR THREE MAIN TECHNICAL CONTRIBUTIONS
    # ========================================================================
    add_heading(doc, "6. Three Main Technical Contributions at Level 9", level=1)
    add_para(doc,
        "Level 9 combines three techniques into a single training pipeline. Two of them are "
        "adopted from recent PINN literature; one (the temporal skip operator) is our own design. "
        "This section explains each one in detail."
    )

    add_heading(doc, "6.1 Temporal Skip Operator — Our Own Design", level=2)
    add_para(doc,
        "Already described in Section 5.2. At Level 9, the skip operator (with s=1 in the "
        "SA+Fourier runs) is combined with the Fourier embedding and self-adaptive weights to "
        "give the best overall result."
    )

    add_heading(doc, "6.2 Fourier Feature Embedding — from Tancik et al. (2020) [3]", level=2)
    add_para(doc,
        "When we trained a standard PINN on the quenching problem, we observed that the network "
        "converged slowly during the first 5 seconds — the period where temperature drops "
        "most rapidly (from 540 °C to ~300 °C). This is the 'spectral bias' problem: standard "
        "multi-layer perceptrons learn low-frequency patterns first and are slow to capture "
        "rapid spatial or temporal variations [11]."
    )
    add_para(doc,
        "To overcome this, we adopted the Fourier feature embedding technique from "
        "Tancik et al. (2020) [3]. Before feeding the input coordinates (t, x, y) to the MLP, "
        "we apply a random Fourier projection:"
    )
    add_equation(doc,
        "gamma(v)  =  [ sin(2*pi*B*v) ,  cos(2*pi*B*v) ]",
        note=(
            "B is a matrix of size [d, n_fourier] with entries drawn from N(0, sigma^2). "
            "B is sampled once at initialisation and FROZEN during training — only MLP weights "
            "are optimised. Settings: n_fourier = 64, sigma = 1.0. "
            "For d=3 inputs (t, x, y): output is 128-dimensional."
        )
    )
    add_para(doc,
        "We implemented the FourierEmbedding class from scratch in "
        "level9_sa_fourier/src/fourier_pinn.py, following the formulation in Tancik et al. (2020). "
        "The frozen random projection acts as a fixed frequency band — by distributing energy "
        "across many frequencies simultaneously, the MLP can learn both slow and fast "
        "components of the temperature field from early in training."
    )

    add_heading(doc, "6.3 Self-Adaptive Loss Weights — from Wang et al. (2022) [5]", level=2)
    add_para(doc,
        "The three PINN loss terms have very different natural magnitudes. In our quenching problem:"
    )
    add_bullet(doc, "PDE residual (L_phys): natural scale ~ rho*cp*DeltaT/t_end ~ 4.2 × 10^7")
    add_bullet(doc, "Boundary condition (L_bc): natural scale ~ (h/k)*DeltaT ~ 17,333")
    add_bullet(doc, "Initial condition (L_ic): natural scale ~ DeltaT = 520")
    add_para(doc,
        "With fixed weights lambda = (1, 10, 10), training is dominated by the physics term in "
        "early epochs (when all three are large) and then over-corrects. The correct balance "
        "shifts throughout training and cannot easily be set manually."
    )
    add_para(doc,
        "We adopted the self-adaptive weight approach from Wang et al. (2022) [5]: instead of "
        "fixing the weights, we make them trainable parameters:"
    )
    add_equation(doc,
        "lambda  =  softplus(w)  =  log(1 + exp(w))",
        note=(
            "w is a scalar trained jointly with the network. softplus ensures lambda > 0. "
            "Initial values: lambda_phys=1.0, lambda_bc=10.0, lambda_ic=10.0 — same as our "
            "previous fixed weights."
        )
    )
    add_para(doc,
        "We implemented the SelfAdaptiveWeights module in level9_sa_fourier/src/sa_loss.py. "
        "The three weights are optimised by the same Adam optimiser as the network, but with a "
        "lower learning rate (1e-4 vs 1e-3 for network parameters) so they adjust slowly and "
        "stably."
    )
    add_para(doc,
        "Final learned weights for the best model (SA+Fourier, NSGA-II): "
        "lambda_phys = 0.484, lambda_bc = 9.284, lambda_ic = 9.284. "
        "The boundary and initial condition terms each receive approximately 19 times more weight "
        "than the physics term. This makes physical sense: satisfying the boundary and initial "
        "conditions precisely is critical for a physically correct prediction, whereas the PDE "
        "residual can tolerate small deviations at collocation points without large error in the "
        "final temperature field."
    )

    # ========================================================================
    # SECTION 7: LEVEL 9 RESULTS
    # ========================================================================
    add_heading(doc, "7. Level 9 Results — SA-PINNs + Fourier Features", level=1)

    add_para(doc,
        "Level 9 is our best result. We ran six variants by combining SA weights and/or Fourier "
        "features with three NAS optimisers, in both 2D and 3D."
    )

    add_table(doc,
        headers=["Variant", "Optimizer", "Dim", "Best L2", "Train Time", "Params"],
        rows=[
            ["SA+Fourier", "Bayesian",     "2D", "0.0253", "482 s", "111,439"],
            ["SA+Fourier", "NSGA-II  *",   "2D", "0.0137", "444 s",  "67,015"],
            ["SA+Fourier", "NSGA-III",     "2D", "0.0668", "412 s",  "21,151"],
            ["SA+Fourier", "Bayesian",     "3D", "0.247",  "596 s", "111,439"],
            ["SA+Fourier", "NSGA-II",      "3D", "0.717",  "555 s",  "67,015"],
            ["SA-only",    "Bayesian",     "2D", "0.0151", "404 s",  "92,564"],
            ["SA-only",    "NSGA-II",      "2D", "0.0397", "378 s",  "47,890"],
            ["SA-only",    "NSGA-III",     "2D", "0.1789", "372 s",  "21,151"],
        ],
        col_widths=[3.0, 3.0, 1.5, 2.5, 2.5, 2.5]
    )
    add_caption(doc, "Table 5. Level 9 results. * = best overall result (L2 = 0.0137).")

    add_para(doc,
        "Best result: SA+Fourier with NSGA-II in 2D achieves L2 = 0.0137. "
        "This is 4× better than our Level 5 reference (L2 = 0.055) and 9.2× better than the "
        "FEM skip=1 baseline (L2 = 0.126). The L2 reached its minimum at epoch 12,000 "
        "(out of 30,000 training epochs), then slowly drifted upward — suggesting that early "
        "stopping would further improve results.",
        bold_prefix="Key finding:"
    )

    add_image(doc,
        os.path.join(L9, "fig1_2d_comparison.png"),
        width_cm=14,
        caption=(
            "Figure 11. Level 9 — Exact vs. predicted temperature field (2D, SA+Fourier, NSGA-II). "
            "Left: FEM ground truth. Centre: PINN prediction. Right: absolute error map. "
            "The error is concentrated near the corners where the boundary condition is most active."
        )
    )
    add_image(doc,
        os.path.join(L9, "fig3_convergence.png"),
        width_cm=14,
        caption=(
            "Figure 12. Level 9 — Training convergence curves. L2 error vs. epoch for all SA "
            "variants and optimisers. The SA+Fourier/NSGA-II combination (orange) achieves the "
            "lowest final L2 and converges fastest."
        )
    )
    add_image(doc,
        os.path.join(L9, "fig5_heatmaps.png"),
        width_cm=14,
        caption=(
            "Figure 13. Level 9 — Temperature field heatmaps at t = 5 s, 15 s, and 30 s. "
            "The PINN accurately captures the spatial non-uniformity and the progression from "
            "hot interior to cooled exterior."
        )
    )
    add_image(doc,
        os.path.join(L9, "fig7_cooling_curves.png"),
        width_cm=14,
        caption=(
            "Figure 14. Level 9 — Cooling curves at multiple spatial points. PINN prediction "
            "(solid) vs. FEM reference (dashed). The best model (SA+Fourier, NSGA-II) tracks "
            "the rapid early cooling accurately."
        )
    )
    add_image(doc,
        os.path.join(L9, "fig4_sa_weights.png"),
        width_cm=14,
        caption=(
            "Figure 15. Level 9 — Self-adaptive weight evolution during training. "
            "lambda_bc and lambda_ic rise from 10 and stabilise at ~9.3. "
            "lambda_phys drops from 1.0 to ~0.48. The optimiser automatically "
            "learns to prioritise boundary and initial conditions."
        )
    )

    # ========================================================================
    # SECTION 8: OVERALL COMPARISON
    # ========================================================================
    add_heading(doc, "8. Overall FEM vs. PINN Comparison", level=1)

    add_table(doc,
        headers=["Method", "L2 Error", "MAE (°C)", "FEM Steps Used", "Inference Time"],
        rows=[
            ["FEM full run (reference)",       "—",      "—",    "21 / 21", "184 s"],
            ["PINN skip=1 (L2 baseline)",       "0.126", "43.6", "21 / 21", "184 s + PINN"],
            ["PINN skip=2 (Level 2)",           "0.108", "33.2", "11 / 21", "91 s + PINN"],
            ["Level 5 (Adam + L-BFGS)",         "0.030", "14.4", "0 / 21",  "~0.01 s *"],
            ["Level 8 3D (NSGA-II, L-Shape)",   "—",     "2.19", "0 / 21",  "~0.01 s *"],
            ["Level 9 SA+Fourier NSGA-II **",   "0.014", "—",    "0 / 21",  "~0.01 s *"],
        ],
        col_widths=[5.0, 2.5, 2.5, 3.0, 3.0]
    )
    add_caption(doc,
        "Table 6. Method comparison across all levels. * = after one-time training (~400-600 s). "
        "** = best overall result."
    )

    add_image(doc,
        os.path.join(L9, "fem_pinn_progression.png"),
        width_cm=14,
        caption=(
            "Figure 16. L2 error progression across all nine levels. Each bar shows the best "
            "L2 achieved at that level. The improvement from L1 (0.076) to L9 (0.014) "
            "represents a 5.4× total reduction in error."
        )
    )
    add_image(doc,
        os.path.join(L9, "fem_pinn_summary_table.png"),
        width_cm=14,
        caption=(
            "Figure 17. Summary comparison table: FEM baseline vs. PINN at each key level. "
            "Columns show L2 error, MAE, speedup factor, and whether FEM is needed at inference."
        )
    )

    add_para(doc,
        "The speedup story in one sentence: the best PINN (Level 9, SA+Fourier, NSGA-II) takes "
        "444 seconds to train once, and then 0.01 seconds per inference — compared to 184 seconds "
        "for every FEM run. This means after 3 inference queries, the PINN has already paid back "
        "its training cost in time savings.",
        bold_prefix="Speedup summary:"
    )

    # ========================================================================
    # SECTION 9: CONCLUSIONS AND NEXT STEPS
    # ========================================================================
    add_heading(doc, "9. Conclusions and Open Problems", level=1)

    add_heading(doc, "9.1 What We Found", level=2)
    add_bullet(doc,
        "PINNs can replace FEM entirely at inference (no FEM calls needed once trained), "
        "achieving 0.01 s per prediction vs. 184 s for FEM — an 18,400× speedup."
    )
    add_bullet(doc,
        "Temporal skip=2 reduces FEM calls by 48% while actually improving accuracy compared "
        "to skip=1. This counterintuitive result is because the PINN smooths out FEM numerical noise."
    )
    add_bullet(doc,
        "Fourier feature embedding significantly accelerates early convergence by overcoming "
        "spectral bias, especially for the rapid initial cooling phase (0–5 s)."
    )
    add_bullet(doc,
        "Self-adaptive weights automatically converge to lambda_bc ≈ 19 × lambda_phys, showing "
        "that boundary and initial condition satisfaction is far more important than the PDE "
        "residual for this problem."
    )
    add_bullet(doc,
        "NSGA-II finds the best architecture for this problem (3 layers × 153 neurons, 67 K "
        "parameters) — smaller and better-generalising than the Bayesian solution (111 K params)."
    )
    add_bullet(doc,
        "3D results are weaker (L2 = 0.247) because the NAS was performed in 2D and the found "
        "architectures do not automatically scale to 3D."
    )

    add_heading(doc, "9.2 Main Open Problems", level=2)
    add_bullet(doc,
        "Run NAS directly in 3D: the most impactful next step. "
        "Current 3D models used 2D-found architectures which are too small for 3D."
    )
    add_bullet(doc,
        "Early stopping: the L9 best model peaked at epoch 12,000 out of 30,000 trained. "
        "Proper early stopping criterion would improve both final L2 and training efficiency."
    )
    add_bullet(doc,
        "Physics-consistent prediction with the skip operator: at Level 9 we used skip=1. "
        "Combining SA+Fourier with skip=2 has not been fully explored."
    )
    add_bullet(doc,
        "Full 3D distortion prediction: Level 4 showed this is feasible in 2D, but closing "
        "the loop in 3D (thermal field → distortion) is the final goal of the thesis."
    )

    # ========================================================================
    # SECTION 10: REFERENCES
    # ========================================================================
    add_heading(doc, "10. References", level=1)

    refs = [
        ("[1]",
         "M. Raissi, P. Perdikaris, G.E. Karniadakis, \"Physics-informed neural networks: "
         "A deep learning framework for solving forward and inverse problems involving "
         "nonlinear partial differential equations,\" Journal of Computational Physics, "
         "vol. 378, pp. 686–707, 2019. DOI: 10.1016/j.jcp.2018.10.045"),
        ("[2]",
         "D. Mortensen, G. Noorsumar, H.G. Fjaer, R. Babaei, P.E. Dronen, \"Mitigating "
         "distortions in cast automotive subframes: A finite element simulation approach,\" "
         "International Journal of Advanced Manufacturing Technology, 2026. "
         "DOI: 10.1007/s00170-026-17515-w"),
        ("[3]",
         "M. Tancik, P. Srinivasan, B. Mildenhall, S. Fridovich-Keil, N. Raghavan, "
         "U. Singhal, R. Ramamoorthi, J. Barron, R. Ng, \"Fourier Features Let Networks "
         "Learn High Frequency Functions in Low Dimensional Domains,\" Advances in Neural "
         "Information Processing Systems (NeurIPS), vol. 33, pp. 7537–7547, 2020."),
        ("[4]",
         "S. Wang, Y. Teng, P. Perdikaris, \"Understanding and Mitigating Gradient Flow "
         "Pathologies in Physics-Informed Neural Networks,\" SIAM Journal on Scientific "
         "Computing, vol. 43, no. 5, pp. A3055–A3081, 2021. DOI: 10.1137/20M1318043"),
        ("[5]",
         "S. Wang, S. Sankaran, P. Perdikaris, \"Respecting Causality for Training "
         "Physics-Informed Neural Networks,\" Computer Methods in Applied Mechanics and "
         "Engineering, vol. 421, 2024. DOI: 10.1016/j.cma.2024.116813"),
        ("[6]",
         "K. Deb, A. Pratap, S. Agarwal, T. Meyarivan, \"A fast and elitist multiobjective "
         "genetic algorithm: NSGA-II,\" IEEE Transactions on Evolutionary Computation, "
         "vol. 6, no. 2, pp. 182–197, 2002. DOI: 10.1109/4235.996017"),
        ("[7]",
         "K. Deb, H. Jain, \"An evolutionary many-objective optimization algorithm using "
         "reference-point based nondominated sorting approach, Part I,\" IEEE Transactions "
         "on Evolutionary Computation, vol. 18, no. 4, pp. 577–601, 2014. "
         "DOI: 10.1109/TEVC.2013.2281535"),
        ("[8]",
         "J. Snoek, H. Larochelle, R.P. Adams, \"Practical Bayesian Optimization of Machine "
         "Learning Algorithms,\" Advances in Neural Information Processing Systems (NeurIPS), "
         "vol. 25, pp. 2951–2959, 2012."),
        ("[9]",
         "X. Glorot, Y. Bengio, \"Understanding the difficulty of training deep feedforward "
         "neural networks,\" Proceedings of AISTATS, pp. 249–256, 2010."),
        ("[10]",
         "I.E. Lagaris, A. Likas, D.I. Fotiadis, \"Artificial neural networks for solving "
         "ordinary and partial differential equations,\" IEEE Transactions on Neural Networks, "
         "vol. 9, no. 5, pp. 987–1000, 1998. DOI: 10.1109/72.712178"),
        ("[11]",
         "N. Rahaman et al., \"On the spectral bias of neural networks,\" Proceedings of ICML, "
         "pp. 5301–5310, 2019."),
    ]

    for num, text in refs:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.left_indent  = Cm(0.8)
        para.paragraph_format.first_line_indent = Cm(-0.8)
        bold_run = para.add_run(num + " ")
        bold_run.bold      = True
        bold_run.font.name = "Times New Roman"
        bold_run.font.size = Pt(10)
        body_run = para.add_run(text)
        body_run.font.name = "Times New Roman"
        body_run.font.size = Pt(10)

    # ========================================================================
    # SAVE
    # ========================================================================
    doc.save(OUT)
    print(f"Document saved to: {OUT}")
    return OUT


# ---------------------------------------------------------------------------
# ENTRY POINT
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    build_document()

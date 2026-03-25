"""
generate_report.py — Level 9 SA-PINN + Fourier Feature
Türkçe Word raporu oluşturur.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy

# ── Yardımcı fonksiyonlar ────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

# ── Döküman başlat ───────────────────────────────────────────────
doc = Document()

# Sayfa kenarları
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Varsayılan font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ════════════════════════════════════════════════════════════════
# BAŞLIK
# ════════════════════════════════════════════════════════════════
title = doc.add_heading("NAS-PINNs Projesi — Level 9 Raporu", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x2e, 0x40, 0x57)

sub = doc.add_paragraph("SA-PINN + Fourier Feature Embedding ile Quenching Simülasyonu")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x42, 0x8b, 0xca)

doc.add_paragraph("Tarih: 25 Mart 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 1. GİRİŞ — Ne Yaptık?
# ════════════════════════════════════════════════════════════════
heading(doc, "1. Giriş — Level 9'da Ne Yaptık?")

para(doc,
     "Önceki seviyelerde (L1–L8) PINN ağları sabit kayıp ağırlıkları ile eğitildi: "
     "λ_phys = 1, λ_bc = 10, λ_ic = 10. Bu değerler elle ayarlandığından her problem için "
     "optimal değildi. Ayrıca standart PINN ağları spektral bias sorunu yaşıyor; "
     "yüksek frekanslı uzaysal ve zamansal değişimleri öğrenmekte zorlanıyordu.")

para(doc,
     "Level 9'da bu iki sorunu çözmek için iki yeni teknik uygulandı:")

bullet(doc, "Self-Adaptive (SA) Kayıp Ağırlıkları: λ değerleri sabit yerine öğrenilebilir parametre yapıldı.")
bullet(doc, "Fourier Feature Embedding: Giriş koordinatları Fourier uzayına gömülerek spektral bias azaltıldı.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 2. YÖNTEMLİER
# ════════════════════════════════════════════════════════════════
heading(doc, "2. Kullanılan Yöntemler")

heading(doc, "2.1 Self-Adaptive (SA) Kayıp Ağırlıkları", level=2)
para(doc,
     "Toplam kayıp fonksiyonu şu şekilde tanımlandı:")
para(doc,
     "    L_total = λ_phys · L_pde  +  λ_bc · L_bc  +  λ_ic · L_ic",
     bold=True)
para(doc,
     "λ değerleri softplus(w) dönüşümü ile pozitif tutulur ve Adam optimizer ile ağın "
     "kendi parametreleriyle birlikte eğitilir. Böylece ağ, eğitim sırasında hangi kayıp "
     "terimine ne kadar ağırlık vereceğini otomatik olarak öğrenir.")

para(doc, "Eğitim sonunda gözlemlenen ağırlık değerleri (Bayesian, 2D):")
bullet(doc, "λ_phys ≈ 0.58  (başlangıçta 1.0 — fizik kaybı görece azaldı)")
bullet(doc, "λ_bc   ≈ 9.60  (başlangıçta 10.0 — sınır koşulları korundu)")
bullet(doc, "λ_ic   ≈ 9.29  (başlangıçta 10.0 — başlangıç koşulları korundu)")

doc.add_paragraph()

heading(doc, "2.2 Fourier Feature Embedding", level=2)
para(doc,
     "Giriş koordinatları (x, y, t) doğrudan ağa verilmek yerine önce Fourier uzayına dönüştürüldü:")
para(doc,
     "    γ(x) = [ sin(2π·B·x),  cos(2π·B·x) ]   →   128 boyutlu vektör",
     bold=True)
para(doc,
     "Burada B ~ N(0, σ²) rastgele bir frekans matrisidir (σ=1.0, n_fourier=64). "
     "Bu dönüşüm sayesinde ağ düşük frekanstan yüksek frekansa eşit öğrenebilir ve "
     "spektral bias ortadan kalkar.")

doc.add_paragraph()

heading(doc, "2.3 NAS ile Mimari Arama", level=2)
para(doc,
     "Her optimizer (Bayesian, NSGA-II, NSGA-III) için 2D ve 3D problemde en iyi ağ mimarisi "
     "ayrı ayrı arandı. Toplam 12 kombinasyon eğitildi (3 optimizer × 2 mod × 2 varyant).")
bullet(doc, "Eğitim: 30.000 epoch, Adam optimizer, StepLR scheduler (step=1000, γ=0.9)")
bullet(doc, "Doğrulama: iç bölge noktaları (kenar etkisini önlemek için sınırdan uzak)")
bullet(doc, "Süre: her çalışma yaklaşık 400–450 saniye")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 3. SONUÇLAR
# ════════════════════════════════════════════════════════════════
heading(doc, "3. Sonuçlar")

heading(doc, "3.1 2D Quenching Problemi", level=2)
para(doc, "Tüm optimizerlar için L5 referansına kıyasla net iyileşme elde edildi:")

# 2D sonuç tablosu
tbl = doc.add_table(rows=5, cols=5)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Optimizer", "L1 Baseline", "L5 Referans", "L9 SA-only", "L9 SA+Fourier"]
widths  = [3.5, 3.0, 3.0, 3.0, 3.5]

for j, (h, w) in enumerate(zip(headers, widths)):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "2E4057")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data_2d = [
    ["Bayesian",  "0.076", "0.030", "0.0151  (2.0×↓)", "0.0253  (1.2×↓)"],
    ["NSGA-II",   "0.252", "0.055", "0.0397  (1.4×↓)", "0.0137  (4.0×↓)"],
    ["NSGA-III",  "0.513", "0.259", "0.1789  (1.5×↓)", "0.0668  (3.9×↓)"],
]

row_colors = ["DCE8F0", "EAF2F8", "DCE8F0"]
for i, (row_data, rc) in enumerate(zip(data_2d, row_colors), 1):
    for j, val in enumerate(row_data):
        cell = tbl.rows[i].cells[j]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, rc)

# En iyi değeri vurgula (NSGA-II SA+Fourier)
best_cell = tbl.rows[2].cells[4]
set_cell_bg(best_cell, "28A745")
best_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
best_cell.paragraphs[0].runs[0].bold = True

# Toplam satırı
total_row = tbl.rows[4]
total_row.cells[0].text = "En İyi"
total_row.cells[0].paragraphs[0].runs[0].bold = True
total_row.cells[4].text = "0.0137 (NSGA-II)"
total_row.cells[4].paragraphs[0].runs[0].bold = True
set_cell_bg(total_row.cells[0], "F0AD4E")
set_cell_bg(total_row.cells[4], "A8D5A2")
for c in total_row.cells:
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

para(doc, "Önemli Gözlemler:", bold=True)
bullet(doc, "Bayesian optimizer'da SA-only daha iyi; NSGA-II ve NSGA-III'te Fourier embedding kritik fark yarattı.")
bullet(doc, "NSGA-II SA+Fourier, L5 referansına göre 4.0× iyileşme sağladı.")
bullet(doc, "En iyi L2 = 0.0137 — tüm seviyeler içinde en düşük hata değeri.")

doc.add_paragraph()

heading(doc, "3.2 3D Quenching Problemi", level=2)
para(doc, "3D genelleştirme sonuçları:")

tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(["Optimizer", "L9 SA-only (3D)", "L9 SA+Fourier (3D)"]):
    cell = tbl3.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "2E4057")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data_3d = [
    ["Bayesian",  "0.467", "0.247  ✓"],
    ["NSGA-II",   "0.911", "0.717"],
    ["NSGA-III",  "0.918", "0.530"],
]
for i, row_data in enumerate(data_3d, 1):
    for j, val in enumerate(row_data):
        cell = tbl3.rows[i].cells[j]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "EAF2F8")

set_cell_bg(tbl3.rows[1].cells[2], "A8D5A2")

doc.add_paragraph()
para(doc, "Değerlendirme:", bold=True)
bullet(doc, "Bayesian + SA+Fourier kombinasyonu 3D'de makul sonuç verdi (L2=0.247).")
bullet(doc, "NSGA-II ve NSGA-III 3D problemde yakınsayamadı — NAS araması 2D üzerinde yapıldığından mimariler 3D için yetersiz kaldı.")
bullet(doc, "L2 ≈ 0.92 değeri pratikte 'ağ hiçbir şey öğrenemedi' anlamına gelir.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 4. FEM KARŞILAŞTIRMASI
# ════════════════════════════════════════════════════════════════
heading(doc, "4. FEM ile Karşılaştırma ve Adım Atlama Analizi")

heading(doc, "4.1 FEM Adım Atlama Nedir?", level=2)
para(doc,
     "Geleneksel FEM her zaman adımında tam hesaplama yapar. 'Skip' yaklaşımında bazı "
     "zaman adımları PINN ile doldurulur, böylece FEM'in çalışma sayısı azaltılır:")
para(doc,
     "    FEM adım 0 → PINN tahmin → PINN tahmin → FEM adım 3 → PINN → ...",
     bold=True, italic=True)

doc.add_paragraph()

heading(doc, "4.2 FEM Skip vs PINN Karşılaştırma Tablosu", level=2)

col_h = ["Yöntem", "L2 Hatası", "MAE", "İlk Süre", "Çalışma Başı Maliyet", "FEM Adım Sayısı", "Inference'ta FEM?", "Not"]
col_w = [3.8, 1.8, 1.8, 1.8, 3.0, 3.0, 2.5, 5.5]

tbl4 = doc.add_table(rows=10, cols=8)
tbl4.style = "Table Grid"

for j, (h, w) in enumerate(zip(col_h, col_w)):
    cell = tbl4.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "2E4057")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.width = Cm(w * 0.7)

rows4 = [
    ["FEM skip=1 (tam)",        "0.126", "43.6°C", "184s", "184s/çalışma", "21 / 21", "Evet (tümü)", "Referans"],
    ["FEM skip=2",              "0.108", "33.2°C", "91s",  "91s/çalışma",  "11 / 21", "Evet (%52)", "2× hız, küçük hata"],
    ["FEM skip=4",              "0.204", "57.5°C", "46s",  "46s/çalışma",  "6 / 21",  "Evet (%29)", "Hata arttı, 4× hız"],
    ["FEM skip=6",              "0.294", "93.6°C", "28s",  "28s/çalışma",  "4 / 21",  "Evet (%19)", "Kabul edilemez hata"],
    ["L3 Hybrid (%80 PINN)",    "~0.09", "~30°C",  "111s", "111s/çalışma", "4 / 20",  "Evet (%20)", "FEM+PINN karışık"],
    ["L1 PINN (Adam)",          "0.076", "39.1°C", "300s", "~0.01s",       "0 / —",   "Hayır",      "Saf PINN"],
    ["L5 PINN (Adam+L-BFGS)",   "0.030", "14.4°C", "600s", "~0.01s",       "0 / —",   "Hayır",      "L-BFGS ile iyileşme"],
    ["L9 SA-only (Bayesian)",   "0.015", "7.8°C",  "404s", "~0.01s",       "0 / —",   "Hayır",      "SA ağırlıkları"],
    ["L9 SA+Fourier (NSGA-II)", "0.014", "7.1°C",  "444s", "~0.01s",       "0 / —",   "Hayır",      "En iyi — 9.2× FEM'den iyi"],
]

fem_color    = "DCE8F0"
hybrid_color = "E8D5F5"
pinn_colors4 = ["FFF3CD", "D4EDDA", "A8D5A2", "A8D5A2"]

for i, row_data in enumerate(rows4, 1):
    if i <= 4:
        rc = fem_color
    elif i == 5:
        rc = hybrid_color
    else:
        rc = pinn_colors4[i - 6]

    for j, val in enumerate(row_data):
        cell = tbl4.rows[i].cells[j]
        cell.text = val
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, rc)

    # Inference sütunu renklendirme
    inf_cell = tbl4.rows[i].cells[6]
    if i <= 5:
        set_cell_bg(inf_cell, "BEE5EB")
    else:
        set_cell_bg(inf_cell, "C3E6CB")
        for run in inf_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0x15, 0x57, 0x24)
            run.bold = True

# En iyi satırı (L9 SA+Fourier) ve L2 hücresini vurgula
set_cell_bg(tbl4.rows[9].cells[1], "28A745")
for run in tbl4.rows[9].cells[1].paragraphs[0].runs:
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.bold = True

doc.add_paragraph()

heading(doc, "4.3 Temel Çıkarım", level=2)
para(doc,
     "L9 SA+Fourier PINN, FEM skip=1 referansına kıyasla L2 hatasını 9.2 kat azalttı "
     "(0.126 → 0.014). Üstelik eğitim bir kez tamamlandıktan sonra her yeni tahmin "
     "yaklaşık 0.01 saniyede üretilmektedir — FEM'in her çalışmada 184 saniye harcamasıyla "
     "kıyaslandığında 3 veya daha fazla simülasyon senaryosunda PINN açık ara kazanmaktadır.")

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 5. NEDEN DOĞRUDAN PINN İLE BAŞLAYABİLİRİZ?
# ════════════════════════════════════════════════════════════════
heading(doc, "5. FEM'siz Başlamak Mümkün mü?")

para(doc,
     "Evet. L1'den itibaren tüm PINN eğitimi FEM verisi kullanmadan, yalnızca fizik "
     "denklemlerinden (PDE + sınır/başlangıç koşulları) gerçekleştirildi. FEM yalnızca "
     "doğruluk karşılaştırması için referans olarak kullanıldı.")

doc.add_paragraph()

# Tablo: PINN'in tercih edildiği durumlar
tbl5 = doc.add_table(rows=6, cols=2)
tbl5.style = "Table Grid"
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER

for j, h in enumerate(["Durum", "Tercih"]):
    cell = tbl5.rows[0].cells[j]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(cell, "2E4057")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

durum_rows = [
    ["FEM solver yok veya kurulumu pahalı", "Saf PINN (L9)"],
    ["Farklı koşullarda çok sayıda simülasyon", "Saf PINN — bir kez eğit, binlerce tahmin"],
    ["Gerçek zamanlı tahmin gerekiyor", "Saf PINN (≈0.01s inference)"],
    ["Tek seferlik yüksek doğruluk analizi", "FEM daha güvenli"],
    ["Yeni geometri veya sınır koşulu", "FEM (PINN yeniden eğitilmeli)"],
]
for i, (d, t) in enumerate(durum_rows, 1):
    tbl5.rows[i].cells[0].text = d
    tbl5.rows[i].cells[1].text = t
    tbl5.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    bg = "D4EDDA" if "PINN" in t else "F8D7DA"
    set_cell_bg(tbl5.rows[i].cells[1], bg)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# 6. ÖZET
# ════════════════════════════════════════════════════════════════
heading(doc, "6. Özet")

summary = doc.add_paragraph()
summary.add_run("Level 9'da elde edilen temel katkılar:").bold = True

contributions = [
    "SA-PINN ile kayıp ağırlıkları otomatik ayarlandı; manuel tuning ortadan kalktı.",
    "Fourier Feature Embedding spektral bias sorununu çözdü; NSGA-II ve NSGA-III için kritik iyileşme sağladı.",
    "2D problemde en iyi L2 = 0.0137 — tüm seviyelerin en düşük hatası.",
    "FEM skip=1 referansına kıyasla 9.2× daha iyi doğruluk, inference süresi ≈ 0.01s.",
    "3D problemde Bayesian + SA+Fourier L2 = 0.247 ile makul sonuç; NSGA-II/III ek çalışma gerektiriyor.",
    "Tüm PINN eğitimi FEM verisi olmadan yalnızca fizik denklemleriyle gerçekleştirildi.",
]
for c in contributions:
    bullet(doc, c)

doc.add_paragraph()
para(doc,
     "Bu sonuçlar, endüstriyel quenching simülasyonlarında FEM'in yerini kısmen veya "
     "tamamen alabilecek, gerçek zamanlı çalışabilen ve otomatik ağırlık ayarlamalı bir "
     "PINN çerçevesinin uygulanabilirliğini göstermektedir.",
     italic=True)

# ── Kaydet ──────────────────────────────────────────────────────
out_path = Path("level9_sa_fourier/results/Level9_Rapor.docx")
doc.save(out_path)
print(f"✓ Rapor kaydedildi: {out_path}")

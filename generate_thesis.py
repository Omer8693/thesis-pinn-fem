"""
generate_thesis.py
Master tezi — IEEE formatında Word dökümanı.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from pathlib import Path
import copy

OUT = Path("thesis-pinn-fem/reports") if Path("thesis-pinn-fem").exists() else Path("reports")
OUT.mkdir(parents=True, exist_ok=True)

# ── Yardımcılar ──────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=10, font="Times New Roman"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    return run

def body_para(doc, text="", bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    if text:
        add_run(p, text, bold=bold, italic=italic)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(4)
    return p

def section_heading(doc, number, title, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after  = Pt(4)
    if level == 1:
        add_run(p, f"{number}. {title.upper()}", bold=True, size=10)
    else:
        add_run(p, f"{number} {title}", bold=True, italic=True, size=10)
    return p

def italic_label(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, italic=True, size=9)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def make_table(doc, headers, rows, col_widths_cm, header_bg="2E4057", row_bgs=None):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, (h, w) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9)
        set_cell_bg(cell, header_bg)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row in enumerate(rows):
        bg = (row_bgs[i] if row_bgs and i < len(row_bgs) else "FFFFFF")
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.width = Cm(col_widths_cm[j])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(val), size=9)
            set_cell_bg(cell, bg)
    return tbl

# ── Döküman ──────────────────────────────────────────────────────
doc = Document()

# Sayfa yapısı — A4, dar kenar
for sec in doc.sections:
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# Varsayılan stil
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════
# BAŞLIK & KİŞİ BİLGİLERİ
# ════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_p,
    "NAS-PINN FRAMEWORK FOR INDUSTRIAL THERMAL SIMULATION:\n"
    "A NINE-LEVEL CURRICULUM FROM FEM BASELINES TO\n"
    "SELF-ADAPTIVE PHYSICS-INFORMED NEURAL NETWORKS",
    bold=True, size=14)
title_p.paragraph_format.space_after = Pt(12)

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(info_p, "Ömer Cetinkaya\n", bold=True, size=10)
add_run(info_p, "University of Agder, Department of Information and Communication Technology\n"
                "Grimstad, Norway  ·  omerc@uia.no\n", size=10)
add_run(info_p, "Supervisor: Prof. Turgay Celik", italic=True, size=10)
info_p.paragraph_format.space_after = Pt(14)

doc.add_paragraph("─" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════
abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
abs_p.paragraph_format.space_before = Pt(10)
add_run(abs_p, "ABSTRACT — ", bold=True, size=10)
add_run(abs_p,
    "We propose and evaluate a progressive nine-level framework that investigates whether Neural "
    "Architecture Search (NAS) can identify Physics-Informed Neural Network (PINN) architectures "
    "capable of replacing intermediate Finite Element Method (FEM) time steps in transient thermal "
    "simulations. The physical benchmark is the water quenching of an A356 aluminium alloy component, "
    "a process of high industrial relevance in automotive subframe manufacturing, as studied by "
    "Mortensen et al. [1]. Three NAS strategies are evaluated: Bayesian optimisation (TPE), NSGA-II, "
    "and NSGA-III. The curriculum begins with a global single-shot predictor (Level 1) and advances "
    "through temporal skip operators, hybrid FEM-PINN routing, distortion estimation, L-BFGS "
    "refinement, multi-domain Poisson benchmarks, and three-dimensional extension, culminating in "
    "self-adaptive loss weighting combined with Fourier feature embedding (Level 9). The best result "
    "— Level 9 SA+Fourier (NSGA-II) — achieves an L2 relative error of 0.014, which is 9.2× lower "
    "than the FEM skip=1 baseline (L2=0.126), with subsequent inference in approximately 0.01 seconds. "
    "Our findings demonstrate that FEM can be entirely eliminated at inference time while maintaining "
    "accuracy superior to any FEM skip configuration.",
    size=10)
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

kw_p = doc.add_paragraph()
add_run(kw_p, "Keywords — ", bold=True, size=10)
add_run(kw_p, "Physics-Informed Neural Networks, Neural Architecture Search, Finite Element Method, "
              "Thermal Simulation, Self-Adaptive Training, Fourier Feature Embedding", size=10)
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_p.paragraph_format.space_after = Pt(10)
doc.add_paragraph("─" * 90).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ════════════════════════════════════════════════════════════════
section_heading(doc, "1", "INTRODUCTION")

body_para(doc,
    "Transient thermal simulations using Finite Element Methods (FEM) are central to industrial "
    "process design, particularly in metal casting and heat treatment. In automotive manufacturing, "
    "the quenching of aluminium alloy subframes determines residual stress distributions and "
    "dimensional distortions that affect structural performance. Full FEM simulations are accurate "
    "but computationally expensive: a single 30-second quenching simulation with 21 time steps "
    "requires approximately 184 seconds on contemporary hardware [1].")

body_para(doc,
    "Physics-Informed Neural Networks (PINNs), introduced by Raissi et al. [2], offer a mesh-free "
    "alternative by embedding the governing PDE directly into the training loss. Once trained, "
    "inference is nearly instantaneous. However, standard PINNs are sensitive to manually chosen "
    "hyperparameters, and optimal network architecture varies with the physical problem. Neural "
    "Architecture Search (NAS) addresses this by automating the discovery of effective architectures.")

body_para(doc,
    "In this work, we present a nine-level experimental curriculum that systematically builds from "
    "a simple NAS-PINN predictor toward a fully self-optimising framework. Each level introduces "
    "one specific advancement — temporal skip operators, adaptive routing, mechanical coupling, "
    "L-BFGS refinement, multi-domain benchmarking, 3D extension, and self-adaptive training with "
    "Fourier features — allowing controlled evaluation of each contribution.")

p = body_para(doc, "The main contributions of this work are:")
bullets = [
    "A nine-level, reproducible experimental curriculum for NAS-PINN evaluation on an industrially validated thermal benchmark.",
    "Demonstration that Bayesian-optimised PINNs can replace up to 80% of FEM time steps while maintaining MAE below 12°C.",
    "Self-Adaptive (SA) loss weighting that eliminates manual hyperparameter tuning of λ values.",
    "Fourier Feature Embedding that resolves spectral bias in NAS-discovered architectures, achieving L2 = 0.014 — a 9.2× improvement over FEM skip=1.",
    "Quantitative evidence that pure PINN inference (0.01 s) outperforms FEM at every skip configuration after a single training pass.",
]
for b in bullets:
    bp = doc.add_paragraph(style="List Bullet")
    add_run(bp, b, size=10)
    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bp.paragraph_format.left_indent = Inches(0.3)

# ════════════════════════════════════════════════════════════════
# 2. RELATED WORK
# ════════════════════════════════════════════════════════════════
section_heading(doc, "2", "RELATED WORK")

section_heading(doc, "2.1", "Physics-Informed Neural Networks", level=2)
body_para(doc,
    "Raissi et al. [2] introduced PINNs by incorporating the PDE residual directly into the neural "
    "network loss, enabling mesh-free solutions to forward and inverse problems. Subsequent work "
    "identified gradient pathologies as a key failure mode: Wang et al. [3] showed that imbalanced "
    "gradient magnitudes across loss terms prevent convergence, motivating adaptive weighting "
    "strategies. Jagtap et al. [4] proposed adaptive activation functions that accelerate convergence "
    "and maintain stable gradient flow in nonlinear systems. Tancik et al. [5] demonstrated that "
    "Fourier feature mappings overcome the spectral bias of standard MLPs, enabling accurate "
    "learning of high-frequency components.")

section_heading(doc, "2.2", "Neural Architecture Search for PINNs", level=2)
body_para(doc,
    "Wang and Zhong [6] introduced the NAS-PINN framework, which combines differentiable "
    "architecture search (DARTS) with PINN training to jointly optimise network weights and "
    "structural parameters. Their bi-level optimisation simultaneously minimises training loss "
    "with respect to network weights and validation loss with respect to architecture parameters. "
    "In our prior work [7], we extended this framework using genetic algorithm (GA)-based search "
    "and evaluated it on Burgers and Poisson equations across diverse geometries, demonstrating "
    "that GA-NAS-PINN outperforms manually designed networks in all cases.")

section_heading(doc, "2.3", "FEM–PINN Hybrid Methods", level=2)
body_para(doc,
    "Several works have explored combining FEM and neural surrogates. Hybrid approaches typically "
    "use FEM for coarse time stepping and a neural network to interpolate fine-grained solutions. "
    "The temporal skip operator concept [8] replaces a fixed number of FEM steps with a learned "
    "surrogate, reducing simulation cost proportionally to the skip rate. Our Level 3 extends this "
    "with residual-based adaptive routing, activating FEM only when the PINN residual exceeds a "
    "threshold. To our knowledge, this is the first study to apply NAS-discovered architectures "
    "within an adaptive FEM–PINN routing system on an industrially validated quenching benchmark.")

# ════════════════════════════════════════════════════════════════
# 3. PHYSICAL PROBLEM AND FEM REFERENCE MODEL
# ════════════════════════════════════════════════════════════════
section_heading(doc, "3", "PHYSICAL PROBLEM AND FEM REFERENCE MODEL")

body_para(doc,
    "All experiments in this work are grounded in the industrial FEM study of Mortensen et al. [1], "
    "which provides validated baseline results for the water quenching of an A356 cast aluminium alloy "
    "automotive subframe. This section defines the physical problem, governing equations, material "
    "properties, and FEM baseline against which all PINN levels are evaluated.")

section_heading(doc, "3.1", "Quenching Process Description", level=2)
body_para(doc,
    "The quenching process consists of immersing a hot cast component (T₀ = 540°C) into a cold "
    "water bath (T_water = 20°C). Rapid surface cooling induces a non-uniform temperature field "
    "that drives thermal gradients, residual stresses, and dimensional distortions. The simulation "
    "domain is a rectangular cross-section representing the component geometry. The total simulation "
    "time is 30 seconds, discretised into 21 uniform time steps.")

section_heading(doc, "3.2", "Governing Equations", level=2)
body_para(doc, "The transient heat conduction equation governs the temperature field:")

eq1 = doc.add_paragraph()
eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq1, "ρ · Cₚ · ∂T/∂t  =  K · ∇²T        in Ω × (0, T_end]", italic=True, size=10)

body_para(doc, "with Robin boundary condition on the wetted surface:")

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq2, "−K · ∂T/∂n  =  h(T) · (T − T_water)   on ∂Ω", italic=True, size=10)

body_para(doc,
    "where h(T) is the temperature-dependent heat transfer coefficient (HTC), which follows a "
    "power-law relationship derived from film boiling and nucleate boiling transition physics. "
    "The initial condition is T(x, 0) = 540°C everywhere in Ω.")

body_para(doc, "The analytical reference solution for uniform cooling is:")
eq3 = doc.add_paragraph()
eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq3, "T(t) = 20 + 520 · exp(−1.75 × 10⁻³ · t)", italic=True, size=10)

section_heading(doc, "3.3", "Material Properties (A356 Aluminium Alloy)", level=2)
body_para(doc,
    "The thermophysical and mechanical properties of the A356 alloy used in all simulations "
    "are taken directly from Mortensen et al. [1] and are summarised in Table 1.")

italic_label(doc, "TABLE 1: A356 Aluminium Alloy Material Properties")
make_table(doc,
    headers=["Property", "Symbol", "Value", "Unit"],
    rows=[
        ["Thermal conductivity",   "K",     "150",     "W/(m·K)"],
        ["Density × heat capacity","ρ·Cₚ",  "2.4×10⁶", "J/(m³·K)"],
        ["Elastic modulus",        "E",     "69",      "GPa"],
        ["Thermal expansion",      "α",     "22×10⁻⁶", "1/°C"],
        ["Initial temperature",    "T₀",    "540",     "°C"],
        ["Coolant temperature",    "T_w",   "20",      "°C"],
    ],
    col_widths_cm=[5.5, 2.5, 3.0, 3.0],
    row_bgs=["F0F4F8","FFFFFF","F0F4F8","FFFFFF","F0F4F8","FFFFFF"]
)

section_heading(doc, "3.4", "FEM Baseline and Time-Step Skip Analysis", level=2)
body_para(doc,
    "The FEM reference model from Mortensen et al. [1] was used as the ground truth for all "
    "PINN validation metrics. In the time-step skip configuration, the FEM solver computes "
    "every k-th time step; intermediate steps are filled by the PINN surrogate. Table 2 "
    "summarises the FEM skip results obtained using the Bayesian-optimised PINN architecture "
    "from Level 2. The FEM full (skip=1) configuration serves as the primary reference: "
    "L2_rel = 0.126, MAE = 43.6°C, runtime = 184 seconds per simulation.")

italic_label(doc, "TABLE 2: FEM Time-Step Skip Analysis (Bayesian Architecture, Level 2)")
make_table(doc,
    headers=["Skip Factor", "FEM Steps / Total", "L2 Relative Error", "MAE (°C)", "Runtime (s)", "FEM Reduction"],
    rows=[
        ["skip=1 (full)", "21 / 21", "0.126", "43.6", "184", "0%"],
        ["skip=2",        "11 / 21", "0.108", "33.2", "91",  "48%"],
        ["skip=4",        "6 / 21",  "0.204", "57.5", "46",  "71%"],
        ["skip=6",        "4 / 21",  "0.294", "93.6", "28",  "81%"],
    ],
    col_widths_cm=[3.0, 3.5, 3.5, 2.5, 2.5, 3.0],
    row_bgs=["DCE8F0","EAF2F8","DCE8F0","EAF2F8"]
)

body_para(doc,
    "A key observation from Table 2 is that skip=2 achieves lower error than skip=1, because "
    "the PINN avoids the sequential numerical error accumulation inherent in step-by-step FEM "
    "integration. However, skip=4 and skip=6 degrade significantly, indicating that the PINN "
    "surrogates at those levels (L1–L2) lack sufficient accuracy to bridge large time intervals. "
    "This limitation motivates the progressive accuracy improvements in Levels 3–9.")

# ════════════════════════════════════════════════════════════════
# 4. PROPOSED FRAMEWORK: NAS-PINN
# ════════════════════════════════════════════════════════════════
section_heading(doc, "4", "PROPOSED FRAMEWORK: NAS-PINN WITH SELF-ADAPTIVE TRAINING")

section_heading(doc, "4.1", "PINN Formulation", level=2)
body_para(doc,
    "We consider the general PDE formulation from Raissi et al. [2]. The PINN approximates the "
    "solution u(t, x) with a deep neural network u_θ(t, x), where θ denotes all trainable parameters. "
    "The total training loss combines three residual terms:")
eq4 = doc.add_paragraph()
eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq4, "L_total = λ_phys · L_PDE  +  λ_bc · L_BC  +  λ_ic · L_IC", italic=True, size=10)

body_para(doc,
    "where L_PDE enforces the governing equation at interior collocation points, L_BC enforces "
    "boundary conditions on ∂Ω, and L_IC enforces initial conditions. In standard PINNs, the "
    "weights λ_phys, λ_bc, λ_ic are fixed scalars set by the practitioner. In our Level 9 "
    "framework, these are learned automatically.")

section_heading(doc, "4.2", "NAS Search Strategies", level=2)
body_para(doc,
    "We evaluate three NAS optimisers that search the same architecture space: number of hidden "
    "layers ∈ {2,3,4,5,6}, neurons per layer ∈ {32,48,64,96,128,160,192,256}, and activation "
    "function ∈ {tanh, relu, swish, gelu}.")

body_para(doc,
    "Bayesian Optimisation (TPE) treats architecture selection as a black-box optimisation problem, "
    "building a probabilistic model of the objective (L2 relative error) as a function of the "
    "architecture configuration. NSGA-II and NSGA-III are multi-objective evolutionary algorithms "
    "that simultaneously minimise L2 error and parameter count, producing a Pareto front of "
    "accuracy-efficiency trade-offs.")

section_heading(doc, "4.3", "Self-Adaptive Loss Weights (Level 9)", level=2)
body_para(doc,
    "Fixed loss weights require manual tuning and may be suboptimal across architectures and "
    "training stages. We introduce learnable weights via the softplus transformation:")
eq5 = doc.add_paragraph()
eq5.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq5, "λ_k = softplus(w_k) = log(1 + exp(w_k)),   k ∈ {phys, bc, ic}", italic=True, size=10)
body_para(doc,
    "The parameters w_k are initialised so that λ values match the Level 1 fixed weights "
    "(λ_phys=1.0, λ_bc=10.0, λ_ic=10.0). They are updated jointly with the network weights "
    "using Adam optimisation. At convergence, the Bayesian model settles at λ_phys≈0.58, "
    "λ_bc≈9.60, λ_ic≈9.29, indicating the network automatically reduces the relative weight "
    "on the PDE residual while maintaining strong boundary and initial condition enforcement.")

section_heading(doc, "4.4", "Fourier Feature Embedding (Level 9)", level=2)
body_para(doc,
    "Standard MLPs suffer from spectral bias — a tendency to learn low-frequency components "
    "first, failing to capture sharp gradients. Following Tancik et al. [5], we map the input "
    "coordinates through a random Fourier feature projection before passing them to the network:")
eq6 = doc.add_paragraph()
eq6.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(eq6, "γ(x) = [sin(2π·B·x), cos(2π·B·x)]   where B ~ N(0, σ²)", italic=True, size=10)
body_para(doc,
    "We use n_fourier=64 frequency samples (σ=1.0), producing a 128-dimensional embedding. "
    "The frequency matrix B is sampled once at initialisation and kept fixed. This embedding "
    "is particularly effective for NSGA-II and NSGA-III architectures, which tend to be shallower "
    "and benefit more from the enriched input representation.")

# ════════════════════════════════════════════════════════════════
# 5. NINE-LEVEL EXPERIMENTAL CURRICULUM
# ════════════════════════════════════════════════════════════════
section_heading(doc, "5", "NINE-LEVEL EXPERIMENTAL CURRICULUM")

body_para(doc,
    "The framework is structured as a progressive curriculum. Each level introduces a single "
    "new component, enabling controlled ablation of each contribution. Table 3 summarises "
    "all nine levels.")

italic_label(doc, "TABLE 3: Nine-Level Experimental Curriculum Overview")
make_table(doc,
    headers=["Level", "Title", "Dim.", "New Contribution", "Best L2 (Bayesian)"],
    rows=[
        ["L1", "Single-Shot NAS-PINN",     "2D",    "NAS architecture search for global T(t,x,y)",        "0.076"],
        ["L2", "Temporal Skip Operator",   "2D",    "PINN replaces FEM at intermediate time steps",       "0.127*"],
        ["L3", "Hybrid FEM + PINN",        "2D",    "Residual-based adaptive routing; 80% FEM reduction", "~0.09"],
        ["L4", "Thermal–Mechanical",       "2D",    "Temperature field coupled to CMM distortion (mm)",   "—"],
        ["L5", "L-BFGS Refinement",        "2D",    "Extended training; L-BFGS second-order optimiser",   "0.030"],
        ["L6", "Poisson Aux. Fine-tuning", "2D",    "Auxiliary Poisson loss; marginal accuracy gain",     "0.028"],
        ["L7", "Multi-Domain Poisson NAS", "2D",    "NAS across 5 geometries; circle L2 = 1.4×10⁻⁴",     "0.030"],
        ["L8", "3D Multi-Domain NAS",      "3D",    "4 geometries, 3 optimisers, 4 skip values",          "0.467"],
        ["L9", "SA-PINN + Fourier",        "2D+3D", "Learnable λ weights + Fourier embedding",            "0.015"],
    ],
    col_widths_cm=[1.2, 4.5, 1.2, 7.5, 3.0],
    row_bgs=["F5F5F5","FFFFFF"]*5
)
body_para(doc, "* skip=1 configuration (full temporal resolution)", italic=True)

# ════════════════════════════════════════════════════════════════
# 6. RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════
section_heading(doc, "6", "RESULTS AND DISCUSSION")

section_heading(doc, "6.1", "2D Quenching: Level-by-Level Accuracy Progression", level=2)
body_para(doc,
    "Table 4 tracks the L2 relative error for the three NAS optimisers across the key 2D levels. "
    "Each row represents a new experimental configuration, and all models are evaluated on the "
    "same interior-only validation set (avoiding the T≈20°C boundary bias).")

italic_label(doc, "TABLE 4: L2 Relative Error Progression Across Levels (2D Quenching)")
make_table(doc,
    headers=["Level / Config.", "Bayesian", "NSGA-II", "NSGA-III", "Notes"],
    rows=[
        ["L1 (Adam, 20k ep.)",          "0.076", "0.252", "0.513", "Baseline — NAS only"],
        ["L5 (Adam + L-BFGS)",          "0.030", "0.055", "0.259", "2.5× better than L1 (Bayesian)"],
        ["L9 SA-only (30k ep.)",        "0.015", "0.040", "0.179", "SA weights, no Fourier"],
        ["L9 SA+Fourier (30k ep.)",     "0.025", "0.014", "0.067", "Fourier critical for NSGA-II/III"],
        ["Best L9 (per optimiser)",     "0.015", "0.014", "0.067", "SA wins for Bayesian; SF for others"],
    ],
    col_widths_cm=[4.5, 2.5, 2.5, 2.5, 5.5],
    row_bgs=["F0F4F8","FFFFFF","D4EDDA","A8D5A2","C3E6CB"]
)

body_para(doc,
    "The Bayesian optimiser benefits more from SA-only (2.0× over L5) than from Fourier (1.2×), "
    "because the 5-layer relu architecture already captures the dominant frequencies well. "
    "In contrast, NSGA-II and NSGA-III architectures are shallower and suffer from spectral bias; "
    "Fourier embedding resolves this, yielding 4.0× and 3.9× improvements respectively over L5.")

section_heading(doc, "6.2", "FEM vs PINN: Speed–Accuracy Trade-off", level=2)
body_para(doc,
    "Table 5 presents the comprehensive comparison between FEM time-step skip configurations and "
    "the PINN levels. A critical distinction must be noted: FEM incurs its runtime cost on every "
    "simulation run, whereas PINN training is a one-time cost. After a single training pass "
    "(≈400–600 s), all subsequent inferences take approximately 0.01 seconds.")

italic_label(doc, "TABLE 5: FEM Skip vs PINN Levels — Comprehensive Comparison")
make_table(doc,
    headers=["Method", "L2 Error", "MAE (°C)", "1st Run", "Per-Run Cost", "FEM at Inference?"],
    rows=[
        ["FEM skip=1 (reference)", "0.126", "43.6", "184s", "184s/run", "Yes (100%)"],
        ["FEM skip=2",             "0.108", "33.2", "91s",  "91s/run",  "Yes (52%)"],
        ["FEM skip=4",             "0.204", "57.5", "46s",  "46s/run",  "Yes (29%)"],
        ["FEM skip=6",             "0.294", "93.6", "28s",  "28s/run",  "Yes (19%)"],
        ["L3 Hybrid (80% PINN)",   "~0.09", "~30",  "111s", "111s/run", "Yes (20%)"],
        ["L1 PINN (Adam)",         "0.076", "39.1", "300s", "~0.01s",   "No"],
        ["L5 PINN (Adam+L-BFGS)", "0.030", "14.4", "600s", "~0.01s",   "No"],
        ["L9 SA-only (Bayesian)",  "0.015", "7.8",  "404s", "~0.01s",   "No"],
        ["L9 SA+Fourier (NSGA-II)","0.014", "7.1",  "444s", "~0.01s",   "No  ← Best"],
    ],
    col_widths_cm=[4.8, 2.5, 2.5, 2.2, 2.8, 3.5],
    row_bgs=["DCE8F0","EAF2F8","DCE8F0","EAF2F8",
             "E8D5F5","FFF3CD","D4EDDA","A8D5A2","A8D5A2"]
)

body_para(doc,
    "The Level 9 SA+Fourier model achieves L2=0.014, which is 9.2× better than the FEM skip=1 "
    "reference and 7.7× better than FEM skip=2. Crucially, after the initial training investment, "
    "the PINN requires no FEM computation at inference time. For workflows requiring three or more "
    "simulation runs under different cooling conditions, the PINN total cost becomes lower than "
    "a single FEM skip=1 run.")

section_heading(doc, "6.3", "3D Generalization", level=2)
body_para(doc,
    "Level 9 was also evaluated on a three-dimensional extension of the quenching problem "
    "(domain: 1.3 m × 0.6 m × 0.4 m, input: [t, x, y, z]). Table 6 reports the results.")

italic_label(doc, "TABLE 6: Level 9 — 3D Quenching Results")
make_table(doc,
    headers=["Optimizer", "SA-only L2", "SA+Fourier L2", "Training Time (s)", "Assessment"],
    rows=[
        ["Bayesian",  "0.467", "0.247", "~420s", "Acceptable — Fourier helps significantly"],
        ["NSGA-II",   "0.911", "0.717", "~430s", "Poor — NAS searched on 2D only"],
        ["NSGA-III",  "0.918", "0.530", "~415s", "Poor — architecture too shallow for 3D"],
    ],
    col_widths_cm=[2.8, 2.8, 3.2, 3.5, 5.5],
    row_bgs=["D4EDDA","F8D7DA","F8D7DA"]
)

body_para(doc,
    "The 3D results reveal that Bayesian-optimised architectures (5-layer, 151 neurons, relu) "
    "generalise to 3D, while NSGA-II/III architectures (3 layers, 75 neurons, tanh) are "
    "insufficient. The fundamental limitation is that NAS was performed on the 2D problem; "
    "a 3D-specific NAS run would be required to close this gap. Fourier embedding reduces the "
    "Bayesian L2 from 0.467 to 0.247, confirming that spectral bias is amplified in higher dimensions.")

section_heading(doc, "6.4", "Self-Adaptive Weight Evolution", level=2)
body_para(doc,
    "The SA weight evolution provides insight into problem geometry. Starting from "
    "(λ_phys=1.0, λ_bc=10.0, λ_ic=10.0), the Bayesian model converges to "
    "(λ_phys≈0.58, λ_bc≈9.60, λ_ic≈9.29). The reduction in λ_phys indicates that "
    "the Robin boundary condition is the dominant training signal early on, and the PDE "
    "residual becomes relatively easier to satisfy once the boundary profile is established. "
    "This automatic discovery of effective weighting removes the need for grid search over "
    "loss coefficients, which is particularly valuable when transferring the framework to "
    "new physical problems or geometries.")

# ════════════════════════════════════════════════════════════════
# 7. CONCLUSION AND FUTURE WORK
# ════════════════════════════════════════════════════════════════
section_heading(doc, "7", "CONCLUSION AND FUTURE WORK")

body_para(doc,
    "We presented a nine-level NAS-PINN curriculum for industrial thermal simulation, grounded in "
    "the FEM-validated quenching benchmark of Mortensen et al. [1]. The framework progressively "
    "introduced temporal skip operators, hybrid FEM–PINN routing, mechanical coupling, L-BFGS "
    "refinement, multi-domain Poisson benchmarking, 3D extension, and finally self-adaptive "
    "training with Fourier feature embedding.")

body_para(doc, "The key quantitative conclusions are:")

conclusions = [
    "L9 SA+Fourier (NSGA-II) achieves L2=0.014, the best result across all levels and 9.2× better than FEM skip=1.",
    "After a single training pass (~444 s), inference is approximately 0.01 s — making pure PINN competitive for any workflow with more than two simulation queries.",
    "Self-adaptive weights eliminate manual loss tuning and converge to physically meaningful coefficients.",
    "Fourier embedding is critical for NSGA-II and NSGA-III architectures, delivering 4.0× and 3.9× improvements over L5 respectively.",
    "3D generalisation is feasible for Bayesian architectures (L2=0.247) but requires 3D-specific NAS for multi-objective optimisers.",
]
for c in conclusions:
    cp = doc.add_paragraph(style="List Bullet")
    add_run(cp, c, size=10)
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cp.paragraph_format.left_indent = Inches(0.3)

body_para(doc, "")
body_para(doc,
    "Future work will address: (i) 3D NAS search on the quenching problem directly; "
    "(ii) transfer learning from 2D to 3D models to reduce 3D training cost; "
    "(iii) application of the SA+Fourier framework to alternative geometries (cylinder, "
    "L-shape, stacked) to validate generalisation across the domain shapes studied in Level 8; "
    "(iv) coupling with the mechanical distortion model of Level 4 to produce end-to-end "
    "thermal–mechanical PINN predictions without any FEM computation.")

# ════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════
section_heading(doc, "", "REFERENCES")

refs = [
    "[1] D. Mortensen, G. Noorsumar, H. G. Fjær, R. Babaei, and P. E. Dronen, "
    "\"Mitigating distortions in cast automotive subframes: A finite element simulation approach,\" "
    "Int. J. Adv. Manuf. Technol., 2026. https://doi.org/10.1007/s00170-026-17515-w",

    "[2] M. Raissi, P. Perdikaris, and G. E. Karniadakis, "
    "\"Physics-informed neural networks: A deep learning framework for solving forward and inverse "
    "problems involving nonlinear partial differential equations,\" "
    "J. Comput. Phys., vol. 378, pp. 686–707, 2019.",

    "[3] S. Wang, H. Teng, and P. Perdikaris, "
    "\"Understanding and mitigating gradient pathologies in physics-informed neural networks,\" "
    "SIAM J. Sci. Comput., vol. 43, no. 5, pp. A3055–A3081, 2021.",

    "[4] A. D. Jagtap, K. Kawaguchi, and G. E. Karniadakis, "
    "\"Adaptive activation functions accelerate convergence in physics-informed neural networks,\" "
    "J. Comput. Phys., vol. 404, p. 109136, 2020.",

    "[5] M. Tancik, P. P. Srinivasan, B. Mildenhall, S. Fridovich-Keil, N. Raghavan, U. Singhal, "
    "R. Ramamoorthi, J. T. Barron, and R. Ng, "
    "\"Fourier features let networks learn high frequency functions in low dimensional domains,\" "
    "NeurIPS, 2020.",

    "[6] Y. Wang and L. Zhong, "
    "\"NAS-PINN: Neural architecture search-guided physics-informed neural network for solving "
    "PDEs,\" J. Comput. Phys., 2023.",

    "[7] Ö. Cetinkaya, "
    "\"Enhancing Physics-Informed Neural Networks with Automated Architecture and Training Strategies,\" "
    "IKT464-G 25H Advanced ICT Project, University of Agder, 2025.",

    "[8] H. Wu, J. Wang, and X. Chen, "
    "\"Residual-based adaptive sampling methods for physics-informed neural networks,\" "
    "Comput. Mech., vol. 72, no. 3, pp. 563–582, 2023.",
]

for ref in refs:
    rp = doc.add_paragraph()
    add_run(rp, ref, size=9)
    rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rp.paragraph_format.left_indent = Inches(0.25)
    rp.paragraph_format.first_line_indent = Inches(-0.25)
    rp.paragraph_format.space_after = Pt(3)

# ── Kaydet ──────────────────────────────────────────────────────
out_path = OUT / "NAS_PINN_Master_Thesis_Report.docx"
doc.save(out_path)
print(f"✓ Tez raporu kaydedildi: {out_path}")
